import docx

doc = docx.Document("BC_TravelMemoryMap_Final (1).docx")
print("Total paragraphs:", len(doc.paragraphs))
print("Total tables:", len(doc.tables))

# Print first 100 paragraphs with their style to find section 2.2 and 3.2
import sys
sys.stdout.reconfigure(encoding='utf-8')

for i, p in enumerate(doc.paragraphs[:200]):
    if p.text.strip():
        print(f"[{i}] ({p.style.name}): {p.text[:100]}")
