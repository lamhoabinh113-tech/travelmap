import os
import sys
import subprocess

# Check and install python-docx if not present
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("50 CÂU HỎI BẢO VỆ ĐỀ TÀI - TRAVEL MEMORY MAP")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 82, 118)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Tài liệu hướng dẫn câu trả lời chi tiết và ôn tập bảo vệ đồ án")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("-" * 80)
    
    questions = [
        # Part 1
        ("PHẦN 1: GIỚI THIỆU ĐỀ TÀI & TỔNG QUAN HỆ THỐNG", [
            ("1. Em hãy giới thiệu ngắn gọn về đề tài của mình?", 
             "Đề tài của em là Travel Memory Map (Bản đồ kỷ niệm du lịch). Đây là một ứng dụng web xây dựng trên nền tảng PHP theo mô hình MVC, kết hợp thư viện bản đồ số Leaflet JS. Ứng dụng cho phép người dùng lưu trữ, định vị và trực quan hóa các địa điểm đã đi qua dưới dạng các điểm đánh dấu (marker) trên bản đồ trực quan, đi kèm với hình ảnh kỷ niệm, ngày ghé thăm, ghi chú cảm xúc và khả năng chia sẻ, tương tác bạn bè."),
            
            ("2. Mục tiêu chính của hệ thống là gì?",
             "Mục tiêu chính của hệ thống là giúp người dùng cá nhân và nhóm bạn bè:\n- Lưu trữ và số hóa nhật ký hành trình du lịch gắn với tọa độ địa lý cụ thể (kinh độ/vĩ độ).\n- Tạo và quản lý các chuyến đi (Trips) chung, cho phép nhiều thành viên cùng tham gia đóng góp điểm đến.\n- Tăng tính tương tác thông qua tính năng mạng xã hội (kết bạn, bình luận, bày tỏ cảm xúc trên các kỷ niệm).\n- Tích hợp trợ lý AI thông minh để gợi ý lịch trình, tư vấn địa điểm du lịch dựa trên dữ liệu người dùng."),
            
            ("3. Hệ thống giải quyết vấn đề gì trong thực tế?",
             "Vấn đề thực tế: Thông thường, người dùng lưu trữ ảnh du lịch rời rạc trên điện thoại hoặc đăng tải trên mạng xã hội lớn (Facebook, Instagram). Tuy nhiên, các hình ảnh này không được liên kết trực quan theo lộ trình địa lý, khiến việc xem lại hành trình của một chuyến đi trở nên khó khăn.\nGiải pháp: Hệ thống giải quyết bằng cách liên kết trực tiếp hình ảnh kỷ niệm và cảm xúc với bản đồ số động. Người dùng nhìn vào bản đồ là biết ngay mình đã đi qua những tỉnh thành nào, lộ trình di chuyển ra sao và lưu lại kỷ niệm gì ở đó."),
            
            ("4. Đối tượng sử dụng của hệ thống là ai?",
             "Đối tượng sử dụng bao gồm:\n- Người yêu thích du lịch (Travelers/Phượt thủ): Muốn có một không gian riêng tư hoặc chia sẻ để lưu giữ hành trình.\n- Nhóm bạn bè: Muốn cùng nhau xây dựng album kỷ niệm chung cho những chuyến đi tập thể.\n- Blogger du lịch: Muốn chia sẻ bản đồ hành trình thực tế một cách sinh động tới người theo dõi.")
        ]),
        
        # Part 2
        ("PHẦN 2: KIẾN THỨC PHP CƠ BẢN & CƠ CHẾ HOẠT ĐỘNG", [
            ("5. PHP chạy phía Client hay Server?",
             "PHP chạy ở phía Server (Server-side). Mã nguồn PHP được biên dịch và thực thi trên máy chủ web, sau đó kết quả trả về cho Client (trình duyệt của người dùng) chỉ là mã HTML/CSS/JavaScript. Trình duyệt client không thể nhìn thấy hay chạy trực tiếp code PHP."),
            
            ("6. Quy trình xử lý một trang PHP diễn ra như thế nào?",
             "Quy trình gồm 4 bước chính:\n1. Gửi yêu cầu (Request): Người dùng nhập URL hoặc click vào một liên kết, trình duyệt gửi HTTP Request đến Web Server (ví dụ: Apache/Nginx).\n2. Xử lý tại Server: Web Server nhận yêu cầu, phát hiện đây là file .php nên chuyển tiếp cho PHP Engine xử lý.\n3. Thực thi mã PHP: PHP Engine biên dịch mã PHP, kết nối cơ sở dữ liệu MySQL (nếu cần) để lấy dữ liệu, sau đó kết hợp dữ liệu đó để sinh ra mã HTML/CSS/JS tĩnh.\n4. Phản hồi (Response): Web Server nhận kết quả HTML từ PHP Engine và gửi ngược lại về trình duyệt của Client để hiển thị giao diện cho người dùng."),
            
            ("7. include và require khác nhau như thế nào?",
             "Cả hai đều dùng để nạp một file khác vào file hiện tại, nhưng khác nhau ở cách xử lý lỗi:\n- include: Nếu file cần nạp không tồn tại hoặc bị lỗi, PHP sẽ phát một cảnh báo (Warning) nhưng vẫn tiếp tục thực thi các dòng lệnh phía sau.\n- require: Nếu file cần nạp bị lỗi hoặc không tìm thấy, PHP sẽ phát ra một lỗi nghiêm trọng (Fatal Error) và dừng ngay lập tức toàn bộ chương trình."),
            
            ("8. include_once và require_once dùng để làm gì?",
             "Tương tự như include và require, nhưng hậu tố _once giúp PHP kiểm tra xem file đó đã được nạp trước đó trong phiên chạy này chưa. Nếu file đã được nạp rồi thì PHP sẽ bỏ qua không nạp lại nữa. Điều này giúp tránh lỗi định nghĩa lại lớp (redefine class) hoặc hàm (redefine function)."),
            
            ("9. Biến Session là gì?",
             "Session (phiên làm việc) là một biến siêu toàn cục ($_SESSION) dùng để lưu trữ thông tin của người dùng trên máy chủ (Server) trong suốt phiên làm việc của họ. Mỗi client kết nối sẽ được cấp một mã định danh duy nhất gọi là session_id, mã này được lưu ở Client dưới dạng cookie để Server nhận diện đúng Session của từng người."),
            
            ("10. Cookie là gì?",
             "Cookie là các tệp văn bản nhỏ được máy chủ tạo ra và yêu cầu trình duyệt lưu trữ trực tiếp ở máy khách (Client). Trình duyệt sẽ tự động gửi kèm cookie này lên Server trong mỗi lượt yêu cầu (request) tiếp theo. Cookie thường dùng để ghi nhớ thông tin đăng nhập, tùy chọn ngôn ngữ hoặc các thiết lập cá nhân hóa."),
            
            ("11. Phân biệt Session và Cookie?",
             "Session:\n- Vị trí lưu trữ: Lưu trên máy chủ (Server).\n- Tính bảo mật: Cao (người dùng không thể sửa dữ liệu trực tiếp).\n- Dung lượng: Không giới hạn (tùy thuộc vào bộ nhớ server).\n- Thời hạn tồn tại: Thường mất đi khi đóng trình duyệt (hoặc hết hạn session).\n\nCookie:\n- Vị trí lưu trữ: Lưu trên máy khách (Client - Trình duyệt).\n- Tính bảo mật: Thấp hơn (người dùng có thể xem, chỉnh sửa hoặc xóa).\n- Dung lượng: Giới hạn (thường tối đa 4KB cho mỗi cookie).\n- Thời hạn tồn tại: Có thể thiết lập thời gian sống lâu dài (vài ngày, vài tháng)."),
            
            ("12. Khi nào nên dùng Session thay vì Cookie?",
             "Nên dùng Session khi lưu trữ các thông tin nhạy cảm và quan trọng như: Trạng thái đăng nhập (đã xác thực thành công hay chưa), quyền hạn người dùng (Role), thông tin ví tiền/thanh toán. Vì các dữ liệu này cần được bảo mật tuyệt đối trên server, không để người dùng ở client tự ý chỉnh sửa nhằm thay đổi quyền hạn hoặc gian lận dữ liệu."),
            
            ("13. Hàm isset() dùng để làm gì?",
             "Hàm isset() dùng để kiểm tra xem một biến đã được khai báo/khởi tạo và có giá trị khác NULL hay chưa. Hàm trả về true nếu biến tồn tại và khác NULL, ngược lại trả về false."),
            
            ("14. GET và POST khác nhau như thế nào?",
             "GET: Gửi dữ liệu đính kèm trên thanh địa chỉ URL (dưới dạng query string ?key=value). Dữ liệu bị giới hạn độ dài (khoảng 2048 ký tự), hiển thị công khai nên kém bảo mật, có thể lưu vào lịch sử duyệt web hoặc lưu bookmark.\nPOST: Gửi dữ liệu ẩn trong thân (body) của HTTP request. Không giới hạn dung lượng gửi, dữ liệu không hiện trên URL nên bảo mật hơn, thường dùng khi gửi mật khẩu, upload file hoặc thực hiện các thao tác ghi dữ liệu."),
            
            ("15. Khi nào sử dụng POST thay vì GET?",
             "Sử dụng POST khi:\n- Gửi thông tin nhạy cảm (nhập mật khẩu, thông tin thẻ tín dụng).\n- Thực hiện các hành động làm thay đổi cơ sở dữ liệu (Insert, Update, Delete) để tránh việc người dùng vô tình trigger lại hành động khi reload trang hoặc click link.\n- Khi cần gửi tệp tin (upload file) hoặc các khối dữ liệu có dung lượng lớn."),
            
            ("16. Cách upload file trong PHP?",
             "1. Ở giao diện (HTML Form), phải thiết lập thuộc tính method=\"POST\" và enctype=\"multipart/form-data\".\n2. Trong code PHP, nhận file thông qua biến siêu toàn cục $_FILES['input_name'].\n3. Kiểm tra các thông tin như: Kích thước, định dạng file mở rộng (png, jpg...), mã lỗi.\n4. Sử dụng hàm move_uploaded_file() để di chuyển file từ thư mục tạm thời sang thư mục lưu trữ thực tế trên server."),
            
            ("17. Làm thế nào để kiểm tra dữ liệu đầu vào hợp lệ?",
             "Phải kiểm tra song song ở cả hai phía:\n- Phía Client (Front-end): Dùng thuộc tính HTML5 (required, type=\"email\") và dùng JavaScript để kiểm tra nhanh trước khi gửi form (giúp cải thiện trải nghiệm người dùng, báo lỗi ngay lập tức mà không cần tải lại trang).\n- Phía Server (Back-end - Bắt buộc): Vì người dùng có thể cố tình tắt JS hoặc dùng công cụ như Postman để bỏ qua kiểm tra client. Ở PHP, ta sử dụng các hàm làm sạch như filter_var() (kiểm tra định dạng email, IP), htmlspecialchars(), strip_tags() (loại bỏ mã script độc hại), và ép kiểu dữ liệu rõ ràng."),
            
            ("18. Nếu dữ liệu lên đến hàng triệu bản ghi thì xử lý thế nào?",
             "Cần tối ưu hệ thống theo các cách sau:\n1. Phía Database (MySQL): Tạo chỉ mục (Index) trên các cột thường xuyên dùng để tìm kiếm hoặc liên kết khóa ngoại (như user_id, created_at).\n2. Tối ưu câu lệnh truy vấn: Không sử dụng SELECT * bừa bãi mà chỉ lấy những trường cần thiết. Sử dụng cơ chế phân trang (Pagination) với LIMIT và OFFSET để mỗi lần gọi chỉ lấy một lượng nhỏ dữ liệu (ví dụ: 10-20 dòng) thay vì tải toàn bộ hàng triệu dòng lên RAM.\n3. Caching: Áp dụng các giải pháp lưu trữ bộ nhớ đệm (Redis, Memcached) cho các dữ liệu ít thay đổi để giảm tải truy vấn trực tiếp vào database.\n4. Kiến trúc phần cứng: Thiết lập cơ chế phân mảnh (Sharding/Partitioning) hoặc mô hình Master-Slave (đọc/ghi riêng biệt) nếu hệ thống tiếp tục phình to.")
        ]),
        
        # Part 3
        ("PHẦN 3: KIẾN THỨC VỀ MÔ HÌNH MVC", [
            ("19. MVC là gì?",
             "MVC (Model - View - Controller) là một mẫu kiến trúc phần mềm phổ biến được sử dụng rộng rãi trong phát triển web. Mục đích cốt lõi của nó là phân tách ứng dụng thành 3 thành phần độc lập nhằm giúp mã nguồn rõ ràng, dễ bảo trì, dễ mở rộng và tăng khả năng tái sử dụng code."),
            
            ("20. MVC gồm những thành phần nào?",
             "3 thành phần của MVC bao gồm:\n- M (Model - Mô hình): Quản lý dữ liệu và các quy tắc nghiệp vụ (Business logic).\n- V (View - Giao diện): Hiển thị thông tin trực quan cho người dùng (mã HTML/CSS/JS).\n- C (Controller - Bộ điều khiển): Cầu nối trung gian tiếp nhận yêu cầu, điều hướng và điều phối hoạt động giữa Model và View."),
            
            ("21. Vai trò của Model?",
             "Model chịu trách nhiệm tương tác trực tiếp với cơ sở dữ liệu. Nó thực hiện các câu lệnh SQL (SELECT, INSERT, UPDATE, DELETE) để đọc/ghi dữ liệu, xử lý tính toán logic nghiệp vụ và trả về kết quả thô cho Controller. Trong dự án của em có các Model như UserModel.php, LocationModel.php, TripModel.php."),
            
            ("22. Vai trò của View?",
             "View nhận dữ liệu sạch đã được xử lý từ Controller và có nhiệm vụ trình bày dữ liệu đó ra giao diện đồ họa đẹp mắt dưới dạng HTML, CSS để hiển thị lên trình duyệt cho người dùng cuối tương tác. View không được phép truy cập hay truy vấn trực tiếp vào Database."),
            
            ("23. Vai trò của Controller?",
             "Controller tiếp nhận các yêu cầu (HTTP request) từ phía người dùng gửi lên. Nó sẽ quyết định xem cần gọi Model nào để lấy dữ liệu, xử lý logic gì, sau đó chọn View thích hợp để hiển thị dữ liệu đó ra cho người dùng."),
            
            ("24. Luồng xử lý trong MVC diễn ra như thế nào?",
             "Luồng xử lý tuần tự gồm các bước:\n1. Người dùng tương tác trên giao diện gửi một yêu cầu (Request) kèm URL tới hệ thống.\n2. Router phân tích URL và chuyển yêu cầu đến Controller tương ứng.\n3. Controller tiếp nhận thông tin, gọi Model tương ứng để truy vấn dữ liệu từ MySQL.\n4. Model lấy dữ liệu từ cơ sở dữ liệu, thực hiện tính toán rồi trả kết quả về cho Controller.\n5. Controller nhận dữ liệu từ Model và truyền dữ liệu đó qua View.\n6. View lồng ghép dữ liệu vào mã giao diện HTML rồi render trả lại cho người dùng hiển thị trên trình duyệt."),
            
            ("25. Tại sao phải sử dụng MVC?",
             "Sử dụng MVC giúp dự án:\n- Phân tách rõ ràng trách nhiệm (Separation of Concerns): Lập trình viên Front-end tập trung làm View, lập trình viên Back-end tập trung làm Model/Controller mà không sợ đè code của nhau.\n- Cấu trúc dự án ngăn nắp, không bị tình trạng viết code \"mỳ ống\" (spaghetti code) trộn lẫn cả PHP, HTML và SQL vào cùng một file.\n- Dễ dàng phát hiện và sửa lỗi, bảo trì hệ thống hoặc nâng cấp tính năng mới mà không sợ ảnh hưởng đến các phần khác."),
            
            ("26. MVC có ưu điểm gì so với lập trình PHP thuần?",
             "So với viết PHP thuần (nơi một file chứa tất cả từ kết nối DB, truy vấn SQL, đến viết HTML/CSS):\n- MVC quản lý mã nguồn có tổ chức hơn, giúp mã nguồn tái sử dụng được (ví dụ: một hàm trong Model có thể được gọi bởi nhiều Controller khác nhau).\n- Dễ dàng xây dựng hệ thống định tuyến (Routing) giúp đường dẫn URL đẹp, thân thiện với SEO hơn.\n- Hỗ trợ làm việc nhóm song song cực tốt và dễ bảo trì dự án khi quy mô phình to."),
            
            ("27. Trong dự án của em Controller nào quan trọng nhất?",
             "Controller quan trọng nhất trong dự án của em là LocationController (LocationController.php).\nLý do: Đây là controller cốt lõi điều phối toàn bộ chức năng chính của ứng dụng: Xử lý hiển thị bản đồ tương tác (Dashboard), thêm các điểm kỷ niệm mới (định vị tọa độ vĩ độ/kinh độ), tải và hiển thị danh sách album ảnh kỷ niệm, quản lý quyền riêng tư của từng điểm đi, xử lý tương tác thích (like), bình luận (comment) và đồng bộ trạng thái cảm xúc của người dùng tại các điểm đến."),
            
            ("28. Em đã xây dựng Router như thế nào?",
             "Em xây dựng hệ thống Front Controller với file public/index.php làm điểm tiếp nhận duy nhất:\n1. Sử dụng file cấu hình .htaccess để viết lại đường dẫn (URL Rewriting), chuyển hướng tất cả các request của người dùng (trừ file tĩnh thực tế) về dạng tham số index.php?url=controller/action.\n2. Trong index.php, hệ thống thực hiện phân tích chuỗi tham số từ biến $_GET['url'].\n3. Tách chuỗi này bằng dấu gạch chéo / để xác định tên Controller (phần tử thứ nhất) và Action (phần tử thứ hai).\n4. Kiểm tra xem file Controller có tồn tại trong thư mục app/controllers hay không. Nếu có, thực hiện require_once, khởi tạo đối tượng Controller đó và gọi đến hàm Action tương ứng để xử lý.")
        ]),
        
        # Part 4
        ("PHẦN 4: BẢO MẬT HỆ THỐNG", [
            ("29. Tại sao phải mã hóa mật khẩu? Em sử dụng thuật toán mã hóa nào?",
             "Tại sao: Để bảo vệ tuyệt đối thông tin tài khoản của người dùng. Nếu chẳng may cơ sở dữ liệu của hệ thống bị tấn công rò rỉ (hoặc bị admin xấu cố tình đọc), kẻ xấu cũng chỉ nhìn thấy các chuỗi băm vô nghĩa chứ không biết mật khẩu gốc để đăng nhập.\nThuật toán sử dụng: Em sử dụng thuật toán Bcrypt thông qua hàm password_hash($password, PASSWORD_BCRYPT) có sẵn của PHP."),
            
            ("30. password_hash() và md5() khác nhau như thế nào?",
             "md5() là hàm băm 1 chiều có tốc độ tính toán rất nhanh và luôn tạo ra kết quả giống nhau cho cùng một chuỗi đầu vào (không có muối - salt ngẫu nhiên). Kẻ tấn công có thể dễ dàng giải mã ngược bằng cách tra bảng băm có sẵn (Rainbow Table) hoặc tấn công brute-force thử hàng tỷ mật khẩu mỗi giây.\npassword_hash() (với thuật toán Bcrypt) là thuật toán băm chậm, có độ phức tạp tùy chỉnh (cost). Điểm đặc biệt là nó tự động tạo ra muối (salt) ngẫu nhiên cho mỗi lần băm. Do đó, cùng một mật khẩu băm hai lần sẽ cho ra hai chuỗi kết quả hoàn toàn khác nhau, ngăn chặn triệt để hình thức tra bảng Rainbow Table và làm chậm đáng kể quá trình tấn công brute-force."),
            
            ("31. Vì sao không nên dùng MD5 để lưu mật khẩu?",
             "Vì thuật toán MD5 hiện nay đã lỗi thời, dễ xảy ra lỗi xung đột băm (collision - hai chuỗi khác nhau cho ra cùng một mã hash). Hơn nữa, vì tốc độ băm của nó quá nhanh, kẻ tấn công dùng card đồ họa (GPU) mạnh có thể dò tìm mật khẩu gốc cực kỳ nhanh chóng."),
            
            ("32. Hệ thống dùng phương pháp bảo mật nào?",
             "Hệ thống của em áp dụng các biện pháp bảo mật đa lớp:\n1. Mã hóa mật khẩu: Sử dụng Bcrypt bảo mật cao.\n2. Chống SQL Injection: Sử dụng thư viện PDO với cơ chế Prepared Statements.\n3. Chống tấn công XSS: Sử dụng các hàm htmlspecialchars() và strip_tags() để làm sạch tất cả dữ liệu người dùng nhập trước khi lưu trữ hoặc in ra màn hình.\n4. Bảo mật Session: Thiết lập cơ chế xóa sạch Session khi đăng xuất, cấu hình cookies an toàn.\n5. Phân quyền chặt chẽ: Phân cấp vai trò người dùng (Admin, Moderator, User) kiểm tra quyền trước khi truy cập trang dashboard quản trị hoặc các hành động nhạy cảm."),
            
            ("33. Hệ thống của em chống SQL Injection không và bằng cách nào?",
             "Có. Hệ thống chống SQL Injection bằng cách sử dụng PDO Prepared Statements trong tất cả các truy vấn cơ sở dữ liệu.\nCách thức hoạt động: Thay vì cộng chuỗi trực tiếp dữ liệu người dùng vào câu SQL, hệ thống sẽ biên dịch trước khung của câu lệnh SQL bằng các ký hiệu giữ chỗ (placeholder) như :username. Sau đó, hệ thống liên kết dữ liệu thật thông qua phương thức bindParam() hoặc truyền tham số vào hàm execute(). Lúc này, hệ quản trị cơ sở dữ liệu (MySQL) sẽ hiểu dữ liệu truyền vào chỉ là các tham số thuần túy (dữ liệu thô) chứ không bao giờ biên dịch nó như các lệnh SQL, ngăn chặn hoàn toàn khả năng tiêm mã độc SQL Injection.")
        ]),
        
        # Part 5
        ("PHẦN 5: FRONT-END & THIẾT KẾ GIAO DIỆN (RESPONSIVE)", [
            ("34. Responsive Web Design là gì?",
             "Responsive Web Design (Thiết kế Web đáp ứng) là phương pháp thiết kế giao diện giúp trang web tự động thay đổi bố cục, tỷ lệ hình ảnh và kích thước hiển thị sao cho tối ưu và đẹp mắt trên mọi thiết bị (máy tính để bàn, máy tính bảng, điện thoại di động)."),
            
            ("35. Bootstrap có lợi ích gì?",
             "Cung cấp hệ thống lưới (Grid System) chia màn hình thành 12 cột giúp chia bố cục Responsive cực kỳ nhanh chóng bằng các class tiện ích như col-md-6, col-sm-12.\nCó sẵn thư viện các thành phần giao diện (UI Components) đẹp mắt như Buttons, Modals, Cards, Navbars, Dropdowns giúp lập trình viên không phải viết lại CSS từ đầu.\nĐảm bảo tính tương thích hiển thị tốt trên hầu hết các trình duyệt phổ biến hiện nay."),
            
            ("36. Tại sao giao diện phải Responsive?",
             "Vì lượng người dùng truy cập web thông qua điện thoại thông minh hiện nay chiếm tỷ lệ rất lớn. Nếu giao diện không Responsive, người dùng di động sẽ phải zoom to/nhỏ rất bất tiện, gây ra trải nghiệm người dùng tệ (bad UX). Ngoài ra, Google cũng ưu tiên xếp hạng tìm kiếm cao hơn (SEO) cho các website hỗ trợ Responsive tốt."),
            
            ("37. Em kiểm tra dữ liệu phía Client bằng gì?",
             "Em xử lý kiểm tra dữ liệu phía Client bằng:\n- Các thuộc tính Validation có sẵn của HTML5 (như required, pattern cho mật khẩu, type=\"email\").\n- Kết hợp sử dụng JavaScript để bắt sự kiện gửi form (lắng nghe sự kiện submit), kiểm tra độ khớp mật khẩu, độ dài ký tự và hiển thị thông báo lỗi trực quan ngay dưới ô nhập liệu giúp người dùng sửa thông tin ngay lập tức.")
        ]),
        
        # Part 6
        ("PHẦN 6: HỆ THỐNG MÁY CHỦ, TRIỂN KHAI & MỞ RỘNG (SYSTEM & HOSTING)", [
            ("38. Apache có vai trò gì?",
             "Apache đóng vai trò là phần mềm Web Server. Nó lắng nghe các yêu cầu kết nối từ Client gửi tới cổng 80 (HTTP) hoặc 443 (HTTPS), phân tích yêu cầu đó, phối hợp với PHP Engine để xử lý các tệp tin PHP và gửi trả lại các tài nguyên web (HTML, CSS, hình ảnh) về cho trình duyệt của người dùng hiển thị."),
            
            ("39. Quy trình triển khai website lên Hosting?",
             "Quy trình gồm 6 bước cơ bản:\n1. Chuẩn bị: Nén mã nguồn dự án thành file .zip, và xuất (export) cơ sở dữ liệu MySQL ra file .sql từ phpMyAdmin cục bộ.\n2. Tải mã nguồn: Đăng nhập vào cPanel (hoặc trình quản lý hosting), tải file .zip lên thư mục gốc chạy web công cộng (thường là public_html), sau đó giải nén ra.\n3. Tạo Database: Trên hosting, vào mục quản trị Database tạo một cơ sở dữ liệu trống, tạo một User Database mới và cấp toàn quyền kết nối cho User này vào Database vừa tạo.\n4. Nhập dữ liệu: Vào phpMyAdmin trên hosting, thực hiện Import file .sql đã chuẩn bị vào database mới tạo.\n5. Cập nhật cấu hình: Sửa đổi file cấu hình kết nối database trong code của dự án trên hosting (ví dụ cập nhật thông tin trong file config/db_config.php như Host, Dbname, Username, Password mới tạo trên hosting).\n6. Trỏ tên miền: Cấu hình bản ghi DNS của tên miền (Domain) trỏ về địa chỉ IP của Hosting và tiến hành bật chứng chỉ bảo mật SSL."),
            
            ("40. Domain là gì?",
             "Domain (tên miền) là một địa chỉ định danh bằng chữ thân thiện, dễ nhớ của trang web trên mạng Internet (ví dụ: travelmap.page.gd hoặc google.com). Nó được dùng để thay thế cho địa chỉ IP dạng số của máy chủ vốn rất khó nhớ đối với con người."),
            
            ("41. Hosting là gì?",
             "Hosting là một dịch vụ lưu trữ dữ liệu trực tuyến trên Internet. Nó cung cấp không gian lưu trữ trên các máy chủ có kết nối Internet 24/7 để chứa toàn bộ mã nguồn, hình ảnh, video và cơ sở dữ liệu của trang web, đảm bảo website luôn trực tuyến và bất kỳ ai cũng có thể truy cập được."),
            
            ("42. SSL là gì?",
             "SSL (Secure Sockets Layer) là một công nghệ bảo mật tiêu chuẩn giúp thiết lập một đường truyền mã hóa an toàn giữa máy chủ Web Server và trình duyệt của người dùng (Client). SSL đảm bảo tất cả các dữ liệu truyền đi giữa hai bên không bị nghe lén, đọc trộm hay sửa đổi bởi bên thứ ba."),
            
            ("43. HTTPS khác HTTP như thế nào?",
             "HTTP: Giao thức truyền tải văn bản thuần túy không được mã hóa bảo mật. Nếu tin tặc chặn được gói tin trên đường truyền, họ có thể đọc được nội dung. Cổng kết nối mặc định của HTTP là 80.\nHTTPS: Là phiên bản HTTP bảo mật, được tích hợp thêm chứng chỉ mã hóa dữ liệu SSL/TLS. Toàn bộ dữ liệu truyền tải đều được mã hóa nên cực kỳ an toàn. Cổng kết nối mặc định của HTTPS là 443 và các trang web có HTTPS sẽ có biểu tượng ổ khóa an toàn, được Google đánh giá cao hơn khi làm SEO."),
            
            ("44. Nếu 1000 người truy cập cùng lúc thì hệ thống có vấn đề gì?",
             "Hệ thống có thể gặp các vấn đề quá tải:\n- Server quá tải: CPU và RAM máy chủ tăng vọt, phản hồi rất chậm hoặc gây ra các lỗi nghẽn dịch vụ như 502 Bad Gateway hay 504 Gateway Timeout.\n- Database quá tải: MySQL vượt quá số lượng kết nối đồng thời cho phép (lỗi Too many connections).\nCách khắc phục:\n- Tối ưu hóa các truy vấn SQL và mã nguồn PHP để tiêu tốn ít tài nguyên nhất.\n- Cài đặt công cụ lưu bộ nhớ đệm (caching).\n- Sử dụng mạng phân phối nội dung (CDN như Cloudflare) để gánh tải cho các file tĩnh như hình ảnh kỷ niệm, CSS, JS.\n- Nâng cấp cấu hình gói hosting hoặc sử dụng cơ chế cân bằng tải (Load Balancer) trên cloud server."),
            
            ("45. Nếu database bị mất dữ liệu thì xử lý ra sao?",
             "Hệ thống cần thiết lập quy trình sao lưu dự phòng tự động (Auto Backup) định kỳ (hàng ngày hoặc hàng tuần) thông qua các kịch bản lệnh (cron job) chạy lệnh mysqldump để sao lưu file .sql và lưu trữ ở một máy chủ độc lập hoặc các dịch vụ đám mây an toàn.\nKhi xảy ra sự cố mất mát dữ liệu, quản trị viên sẽ tiến hành nhập lại (Restore) bản sao lưu gần nhất để khôi phục trạng thái hoạt động của website."),
            
            ("46. Nếu muốn phát triển thành ứng dụng di động thì làm thế nào?",
             "Có 2 hướng tiếp cận phổ biến:\n1. Hướng phát triển ứng dụng di động gốc (Native hoặc Hybrid): Ta sẽ giữ nguyên phần Backend hiện tại của dự án nhưng viết thêm các đường dẫn API (RESTful API) bằng PHP để trả về dữ liệu dạng JSON. Sau đó sử dụng các công nghệ làm app di động như Flutter hoặc React Native để xây dựng ứng dụng trên điện thoại và gọi dữ liệu từ API này về hiển thị.\n2. Hướng PWA (Progressive Web App): Cấu hình biến website hiện tại thành một ứng dụng web tiến trình (hiện dự án đã có sẵn tệp cấu hình manifest.json và sw.js để người dùng có thể \"Thêm vào màn hình chính\" trên điện thoại giống như một app thực thụ).")
        ]),
        
        # Part 7
        ("PHẦN 7: TỔNG KẾT & CẢI TIẾN DỰ ÁN", [
            ("47. Nếu được làm lại dự án em sẽ cải tiến những gì?",
             "Em sẽ cải tiến các điểm sau để ứng dụng chuyên nghiệp hơn:\n- Tách biệt hoàn toàn Front-end và Back-end: Sử dụng một framework Front-end hiện đại như React hoặc VueJS để xây dựng giao diện Single Page Application (SPA), giúp trải nghiệm chuyển trang mượt mà không cần load lại trang.\n- Tích hợp dịch vụ lưu trữ đám mây (Cloud Storage): Thay vì lưu ảnh trực tiếp trên ổ cứng server của hosting, em sẽ đẩy ảnh lên các dịch vụ như Cloudinary hoặc Amazon S3 để tối ưu hóa tốc độ tải ảnh và tiết kiệm tài nguyên hosting.\n- Tối ưu hóa bản đồ: Tích hợp tính năng Marker Clustering (nhóm các điểm đánh dấu lại với nhau khi người dùng zoom nhỏ bản đồ) giúp giao diện bản đồ không bị rối mắt khi có quá nhiều địa điểm ở gần nhau."),
            
            ("48. Em học được những gì từ dự án này?",
             "Học được nhiều điều quý giá:\n- Nắm vững kiến trúc phần mềm MVC trong lập trình web PHP thực tế.\n- Kỹ năng làm việc với bản đồ số thông qua Leaflet API và cách làm việc với dữ liệu tọa độ địa lý.\n- Biết cách viết các chức năng bất đồng bộ bằng Ajax/Fetch API để gửi dữ liệu lên server và cập nhật giao diện thời gian thực mà không cần tải lại trang.\n- Kỹ năng bảo mật dữ liệu cơ bản (chống SQL Injection, XSS, hash mật khẩu Bcrypt).\n- Tư duy giải quyết bài toán nghiệp vụ mạng xã hội (tương tác bạn bè, bình luận, chia sẻ quyền riêng tư)."),
            
            ("49. Giải thích hoạt động của MVC trong dự án của em?",
             "Khi người dùng click nút xem Dashboard bản đồ kỷ niệm:\n1. Yêu cầu gửi đi: Trình duyệt gửi request có dạng index.php?url=location/dashboard.\n2. Điều hướng (Router): Router nhận yêu cầu, phân tích chuỗi, tìm thấy Controller là LocationController và gọi phương thức dashboard().\n3. Xử lý dữ liệu (Model): Trong LocationController::dashboard(), hệ thống khởi tạo LocationModel và gọi hàm lấy danh sách địa điểm của người dùng hiện tại từ Database. Model thực thi câu SQL lấy dữ liệu từ các bảng locations và trả về mảng kết quả cho Controller.\n4. Truyền dữ liệu và Hiển thị (View): Controller nhận dữ liệu mảng các địa điểm, sau đó gọi nạp file View tương ứng và truyền mảng địa điểm này vào. View sẽ duyệt mảng dữ liệu này, lồng ghép vào mã HTML/JS của bản đồ Leaflet để vẽ các điểm đánh dấu lên bản đồ hiển thị cho người dùng."),
            
            ("50. PDO (PHP Data Objects) là gì và còn cách nào khác?",
             "PDO là một thư viện hướng đối tượng của PHP đóng vai trò là một lớp giao tiếp đồng nhất giữa ứng dụng PHP và nhiều hệ quản trị cơ sở dữ liệu khác nhau (MySQL, PostgreSQL, SQLite...).\nƯu điểm: Linh hoạt và hỗ trợ cực tốt Prepared Statements để bảo mật chống SQL Injection.\nCách kết nối khác: Sử dụng thư viện MySQLi (MySQL Improved) của PHP. Tuy nhiên, MySQLi chỉ hỗ trợ duy nhất một loại cơ sở dữ liệu là MySQL.")
        ])
    ]
    
    for part_title, part_questions in questions:
        # Add heading for each part
        p_head = doc.add_paragraph()
        run_head = p_head.add_run(part_title)
        run_head.font.name = 'Arial'
        run_head.font.size = Pt(14)
        run_head.font.bold = True
        run_head.font.color.rgb = RGBColor(31, 97, 141)
        p_head.paragraph_format.space_before = Pt(18)
        p_head.paragraph_format.space_after = Pt(6)
        
        for q_title, q_ans in part_questions:
            # Question text
            p_q = doc.add_paragraph()
            run_q = p_q.add_run(q_title)
            run_q.font.name = 'Arial'
            run_q.font.size = Pt(11)
            run_q.font.bold = True
            run_q.font.color.rgb = RGBColor(46, 64, 87)
            p_q.paragraph_format.space_before = Pt(8)
            p_q.paragraph_format.space_after = Pt(2)
            
            # Answer text
            p_a = doc.add_paragraph()
            run_a = p_a.add_run(q_ans)
            run_a.font.name = 'Arial'
            run_a.font.size = Pt(11)
            p_a.paragraph_format.space_after = Pt(10)
            
    # Save to user desktop
    user_profile = os.environ['USERPROFILE']
    desktop_options = [
        os.path.join(user_profile, 'Desktop'),
        os.path.join(user_profile, 'OneDrive', 'Desktop'),
        os.path.join(user_profile, 'OneDrive - HUST', 'Desktop'), # common student accounts
        os.path.join(user_profile, 'OneDrive - Hanoi University of Science and Technology', 'Desktop')
    ]
    
    desktop_path = None
    for path in desktop_options:
        if os.path.exists(path):
            desktop_path = path
            break
            
    if not desktop_path:
        # Fallback to user profile directory if no Desktop folder is found
        desktop_path = user_profile
        
    out_file = os.path.join(desktop_path, 'Cau_Hoi_Bao_Ve_Travel_Map.docx')
    doc.save(out_file)
    print(f"Document saved successfully to {out_file}")

if __name__ == "__main__":
    create_document()
